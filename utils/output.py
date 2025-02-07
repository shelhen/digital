# -*- encoding: utf-8 -*-
"""
--------------------------------------------------------
@File: main.py
@Project: common
@Time: 2024/11/19   01:58
@Author: shelhen
@Email: shelhen@163.com
@Software: PyCharm
--------------------------------------------------------
@Brief:与输出有关的函数，比如输出为图像、excel表格、word内容等。
"""
import os
from pathlib import Path

import pandas as pd
from docx import Document
from docx.oxml import ns
from docx.shared import Pt
import matplotlib.pyplot as plt


def table_translate(matrix: pd.DataFrame, table_name: str, filename: str = "example"):
    filepath = Path(__file__).parent.parent / 'assert/tables' / f"{filename}.docx"
    # 文件若已存在则打开文件, 否则创建新的文档对象
    doc = Document(filepath) if filepath.exists() else Document()
    # 如果包含表格则获取所有表格标题和其索引
    table_map = {table._element.getprevious().text.strip(): idx for idx, table in enumerate(doc.tables)}
    if table_name in table_map:
        # 查找是否有相同标题的表格
        idx = table_map[table_name]
        table = doc.tables[idx]._element
        table.getprevious().getparent().remove(table.getprevious())
        table.getparent().remove(table)
    # 为新表格添加标题
    doc.add_paragraph(f"\n{table_name}")  # 添加表格标题

    # 添加一个表格至Word文档，表格行数为DataFrame的行数加1（为了包括标题），列数与DataFrame的列数相同
    table = doc.add_table(rows=matrix.shape[0]+1, cols=matrix.shape[1])
    table.style = 'Table Grid'
    def set_font_style(_run, en, cn, fontsize=10.5):
        _run.font.name = en  # 英文字体
        _run._element.rPr.rFonts.set(ns.qn('w:eastAsia'), cn)  # 中文字体
        _run.font.size = Pt(fontsize)  # 小五号字体

    for j, col_name in enumerate(matrix.columns):
        cell = table.cell(0, j)
        cell.text = str(col_name)
        set_font_style(cell.paragraphs[0].runs[0],'Times New Roman', '宋体', 10.5)
        cell.paragraphs[0].alignment = 1  # 居中

    # 填充表格数据
    for i, row in matrix.iterrows():
        for j, value in enumerate(row):
            cell = table.cell(i + 1, j)
            cell.text = str(value)
            set_font_style(cell.paragraphs[0].runs[0],'Times New Roman', '宋体', 10.5)
            cell.paragraphs[0].alignment = 1  # 居中
    doc.save(filepath)


def dump_formulation( latex=r"", filename="example", fontsize=12, lamb = 1.8):
    # fontsize=12, # 对应宋体小四
    dpi = 300
    fontsize = fontsize / 72 * dpi
    filepath = Path(__file__).parent / 'assert' / 'formulation' / f"{filename}.svg"
    if filepath.exists():
        os.remove(filepath)
    formula = "${0}$".format(latex.strip())
    fig, ax = plt.subplots(dpi=dpi)
    txt = ax.text(0.5, 0.5, formula, ha='center', va='center', fontsize=fontsize, math_fontfamily='cm')
    ax.axis('off')  # 隐藏坐标轴
    fig.canvas.draw()  # 需要先绘制图形以确保文本被正确测量
    bbox = txt.get_window_extent(renderer=fig.canvas.get_renderer())
    #  # 根据文本大小调整图像大小，让公式长一点
    fig.set_size_inches(bbox.width / fig.dpi * lamb, bbox.height / fig.dpi)
    plt.savefig(filepath, transparent=True, bbox_inches='tight', pad_inches=0)
    # plt.show()


if __name__ == '__main__':
    pass